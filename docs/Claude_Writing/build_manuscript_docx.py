"""
Build BMEiCON_STAGE1_MANUSCRIPT.docx from the clean manuscript content.

IEEE-leaning styling: Times New Roman, single-column (easy to drop into the
two-column IEEE Word template afterward), numbered headings, real tables, and a
numbered reference list. No external deps beyond python-docx.

Run:  python docs/Claude_Writing/build_manuscript_docx.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = Path(__file__).resolve().parent / "BMEiCON_STAGE1_MANUSCRIPT.docx"
IMG_DIR = Path(__file__).resolve().parents[2] / "validation_outputs"
FIG1 = IMG_DIR / "fig_loophole_schematic.png"
FIG2 = IMG_DIR / "fig_feasible_rate_recovery.png"
FIG5 = IMG_DIR / "fig_residual_loophole_schematic.png"
FIG3 = IMG_DIR / "fig_v1_v3_v7_t12_audit_timeseries.png"
FIG4A = IMG_DIR / "fig_so_waveform_canonicalV1_distances.png"
FIG4B = IMG_DIR / "fig_spindle_density_confirm.png"


def add_figure(img_path, width_in=5.5):
    """Center an embedded image; returns the picture paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Inches(width_in))
    p.paragraph_format.space_after = Pt(3)
    return p


def fig_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(10)
    return p

FONT = "Times New Roman"
MONO = "Consolas"

doc = Document()

# ---- base styles -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10)

for h, sz in [("Heading 1", 12), ("Heading 2", 11), ("Title", 18)]:
    st = doc.styles[h]
    st.font.name = FONT
    st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor(0, 0, 0)


def para(text="", *, italic=False, bold=False, align=None, size=10, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.italic = italic
        r.bold = bold
        r.font.name = FONT
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(segments, *, align=None, size=10, space_after=6):
    """segments: list of (text, {bold,italic,mono})."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    for text, opt in segments:
        r = p.add_run(text)
        r.bold = opt.get("bold", False)
        r.italic = opt.get("italic", False)
        r.font.name = MONO if opt.get("mono") else FONT
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def add_table(headers, rows, caption):
    cap = doc.add_paragraph()
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.name = FONT
    cr.font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(3)

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        rr = hdr[i].paragraphs[0].add_run(htxt)
        rr.bold = True
        rr.font.name = FONT
        rr.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val))
            rr.font.name = FONT
            rr.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


# ============================================================================
# TITLE / AUTHOR
# ============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run(
    "Spectrally Similar, Dynamically Invalid: A Single-Subject Diagnosis of "
    "Loophole Exploitation in Neural Mass Model Fitting"
)
tr.bold = True
tr.font.name = FONT
tr.font.size = Pt(18)

para("Author Name(s)", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=0)
para("Affiliation, City, Country  ·  email@domain",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=12)

# ============================================================================
# ABSTRACT
# ============================================================================
abs_p = doc.add_paragraph()
ab = abs_p.add_run("Abstract—")
ab.bold = True
ab.font.name = FONT
ab.font.size = Pt(9)
abs_body = (
    "Personalized neural mass models for sleep EEG are commonly fitted with "
    "spectral objectives. We show that such objectives can select solutions "
    "that are spectrally similar to a target recording yet leave the intended "
    "sleep dynamics unverified—and, when those dynamics are explicitly "
    "audited, invalid—a failure we call shape-alike, spirit-unalike "
    "(形似神不似). Using one Sleep-EDF subject (SC4001, N3) "
    "and a neurolib thalamocortical model (ALN cortical node coupled to a "
    "thalamic node), we diagnose this failure through an iterative V1–V7 "
    "fitting history. A spectral-only baseline reached high residual spectral "
    "similarity (shape_r = 0.8586) but gave no event-level guarantees; auditing "
    "that same solution against 12 event-level criteria shows it passes only "
    "5/12, failing sustained UP states (T3), SO sharpness and regularity "
    "(T4, T6), spindle burstiness (T7), and the entire SO–spindle coupling "
    "triad (T9–T11, modulation index = 0). This makes “spectrally "
    "similar yet dynamically invalid” concrete on a single solution rather "
    "than inferred from version history. Later versions exposed further "
    "optimizer loopholes—persistent-DOWN solutions, SO–spindle "
    "decoupling, fake spindle transients, unreachable spectral spindle rewards. "
    "V7 replaced those rewards with event-level constraints, raising feasible "
    "solutions from 1/4960 (V6) to 364/4960 and yielding a point passing all 12 "
    "constraints. Because reward targets, bounds, duration, and one "
    "spindle-width threshold changed together, we report this trajectory as "
    "confounded and isolate the threshold effect: only 156/4960 V7 points "
    "survive V6’s stricter rule, so roughly half the recovery is threshold "
    "relaxation. Two held-out checks give mixed evidence: V7 is closer to real "
    "EEG on SO waveform morphology (RMSE 0.39 vs 0.57) but produces essentially "
    "zero cortical spindles, where V1 retains some. We claim neither model is "
    "correct; the result is a single-subject diagnostic existence claim and a "
    "transferable loophole-audit workflow."
)
abr = abs_p.add_run(abs_body)
abr.italic = True
abr.font.name = FONT
abr.font.size = Pt(9)

kw = doc.add_paragraph()
kr = kw.add_run("Keywords—")
kr.bold = True
kr.italic = True
kr.font.name = FONT
kr.font.size = Pt(9)
kr2 = kw.add_run(
    "sleep EEG, neural mass model, thalamocortical model, spectral fitting, "
    "slow oscillation, sleep spindle, optimization, specification gaming"
)
kr2.italic = True
kr2.font.name = FONT
kr2.font.size = Pt(9)

# ============================================================================
# I. INTRODUCTION
# ============================================================================
h1("I.  Introduction")
para(
    "Personalized neural mass models are attractive tools for building "
    "mechanistic “digital twins” of brain rhythms. In sleep modeling, "
    "the target signals are often summarized spectrally: N3 sleep is "
    "characterized by slow oscillations, delta activity, and sleep spindles, and "
    "a fitted model is commonly judged by whether its simulated power spectrum "
    "resembles the EEG target [1], [2], [3]. Spectral matching is useful because "
    "it compresses noisy electrophysiological time series into stable "
    "rhythm-level summaries. However, it can also hide mechanistic errors. Two "
    "signals may share similar spectral peaks while differing in whether they "
    "contain real UP/DOWN alternation, waxing–waning spindle events, or "
    "coordinated slow-oscillation–spindle coupling."
)
para(
    "This paper diagnoses that failure mode in one subject: Sleep-EDF SC4001 "
    "during N3 sleep. We call the failure shape-alike, spirit-unalike "
    "(形似神不似): the simulated spectrum resembles the "
    "target, but the underlying dynamics do not implement the intended sleep "
    "physiology. The study is deliberately scoped as a single-subject diagnostic "
    "and existence claim. It does not establish a population-level fitting "
    "method, nor does it claim that the final model is globally more EEG-like on "
    "every metric. Instead, it shows that a standard spectral objective can "
    "select a dynamically invalid solution—one that fails an explicit "
    "event-level audit—and that the iterative history of failed optimizers "
    "is itself evidence of objective loopholes."
)
para(
    "The key framing is to treat the optimizer as an adversarial "
    "loophole-searcher. Differential evolution does not know the intended "
    "physiology; it only searches for parameter settings that satisfy the "
    "objective it is given. If the objective rewards spectral similarity without "
    "event-level safeguards, or if a constraint admits a loophole, the optimizer "
    "can exploit it. Our V1–V7 development history repeatedly exposed such "
    "failures: spectral-only fitting, weak dynamics constraints, PAC-free "
    "spindle solutions, transient events masquerading as spindles, and "
    "FOOOF-based rewards that were unreachable in the SO-stable regime. V7 is the "
    "result of closing those specific loopholes with event-level constraints and "
    "redesigned rewards."
)
para(
    "In estimation terms, this is a statement about objective misspecification "
    "and non-identifiability: a spectral summary does not uniquely determine the "
    "underlying dynamics, so an objective built on it is satisfiable by "
    "dynamically distinct solutions. Our contribution is not that observation, "
    "which is well known, but its iterative operationalization—treating "
    "differential evolution as an adversary and using each feasibility collapse "
    "to localize a specific objective loophole. The principle generalizes to any "
    "simulator fitted to summary statistics; we keep the evidence single-subject "
    "by design."
)
para("The contributions are:", space_after=3)
for i, txt in enumerate([
    "We demonstrate, in a single SC4001 N3 case study, that high spectral "
    "similarity can coexist with dynamically invalid fitted solutions, and we "
    "make this concrete by auditing the V1 spectral-only solution against the "
    "same event-level criteria (Sec. V-A).",
    "We present the V1–V7 fitting history as an empirical loophole audit of "
    "the objective—an adversarial red-teaming protocol—rather than as "
    "an implementation footnote.",
    "We report both the positive event-level result and two independent "
    "held-out checks that give mixed evidence (V7 closer on SO waveform "
    "morphology, V1 closer on cortical spindle density), narrowing the claim to "
    "event-level loophole closure rather than global superiority—and we "
    "surface a residual loophole that survives inside the closed V7 audit.",
], 1):
    doc.add_paragraph(f"{i}) {txt}", style="List Paragraph")

# ============================================================================
# II. RELATED WORK
# ============================================================================
h1("II.  Related Work")
h2("A.  Neural mass fitting for sleep rhythms")
para(
    "Neural mass and mean-field models have been used to reproduce macroscopic "
    "sleep rhythms, including slow oscillations and spindle-band activity [1], "
    "[2], [4]. These models provide mechanistic structure beyond descriptive EEG "
    "statistics: parameters can represent cortical excitability, adaptation, "
    "thalamic conductances, and thalamocortical coupling. Prior work has often "
    "emphasized reproduction of rhythm-level or spectral phenomena. Our work is "
    "complementary but more diagnostic: we ask whether spectral matching is "
    "sufficient as an objective for personalization, and we show that it can "
    "select solutions that satisfy spectral appearance while failing event-level "
    "dynamics."
)
h2("B.  Slow oscillation–spindle physiology")
para(
    "N3 sleep is not defined by a spectrum alone. Slow oscillations involve "
    "alternating DOWN and UP states in cortical activity [5], [6], while spindles "
    "appear as transient waxing–waning sigma-band events associated with "
    "thalamocortical circuitry [7], [8]. Coupling between slow-oscillation phase "
    "and spindle amplitude has been studied as a marker of memory consolidation "
    "and sleep physiology [9], [10], [11]. These phenomena motivate event-level "
    "constraints: a model should not merely have power in the correct bands; it "
    "should produce plausible events with temporal structure. The event metrics "
    "used here are deliberately aligned with established detection conventions "
    "rather than invented: spindle event count and duration follow "
    "amplitude/duration criteria in the spirit of standard detectors [7], [12] "
    "and open implementations such as YASA [13], and the SO criteria follow "
    "zero-crossing/amplitude conventions [5], [6]. The thresholds attached to "
    "these metrics are empirical and subject-specific (Sec. III-D); the metric "
    "forms are not."
)
h2("C.  Optimizer loopholes and specification gaming")
para(
    "The failure studied here resembles specification gaming in machine "
    "learning: an optimizer exploits the literal objective rather than the "
    "designer’s intended goal [14], [15]. In neural model fitting, a "
    "candidate can score well by exploiting spectral artifacts, overly broad "
    "rewards, or incomplete constraints. We import that adversarial-optimizer "
    "framing into personalized sleep modeling. The claim is not that the "
    "optimizer is malicious, but that objective functions define incentives; if "
    "the incentive misses event-level physiology, the search can find "
    "shape-alike but dynamically invalid solutions. Statically, this is the "
    "familiar problem of model/objective misspecification and parameter "
    "non-identifiability. The framing earns its place by converting a static "
    "property into a procedure: each time a patched objective is re-attacked and "
    "feasibility collapses, the collapse localizes the next loophole. The "
    "reusable output is an audit workflow, not the already-accepted conclusion "
    "that constraints help."
)

# ============================================================================
# III. METHODS
# ============================================================================
h1("III.  Methods")
h2("A.  Data and target spectrum")
para(
    "We used Sleep-EDF subject SC4001 and selected N3 sleep epochs. EEG epochs "
    "are aligned to the hypnogram, 30 s N3 epochs are extracted, epochs with "
    "peak-to-peak amplitude above 200 µV are rejected, and Welch PSD "
    "estimates are averaged across accepted epochs (channel EEG Fpz-Cz; 57 "
    "accepted N3 epochs enter the target PSD). The spectral fitting range is "
    "0.5–20 Hz. The held-out SO-waveform check (Sec. III-E) uses a separate, "
    "larger pool of 142 clean N3 epochs under its own QC and is independent of "
    "this target PSD. We use “slow oscillation” (SO) for the "
    "< ~1.25 Hz band that carries the UP/DOWN alternation and “delta” "
    "for the ~1–4 Hz band; the V7 selected point reports T4_freq = 1.25 Hz, "
    "at this SO/delta boundary, and is treated as SO here."
)
para(
    "For the spectral similarity term, the target PSD is decomposed using "
    "FOOOF/specparam-style aperiodic fitting [16]. The comparison is not raw PSD "
    "correlation: both target and simulated PSDs are converted to log-domain "
    "residual spectra after subtracting the fitted aperiodic background, and "
    "shape_r is the Pearson correlation between those residual spectra after "
    "aligning the simulated PSD to the target FOOOF frequency grid. This makes "
    "the objective more rhythm-specific than raw power matching, but the V1 "
    "result shows that even this residual spectral objective does not guarantee "
    "correct dynamics."
)
h2("B.  Thalamocortical model and free parameters")
para(
    "The simulator is a two-node thalamocortical mean-field model implemented "
    "with neurolib’s MultiModel: an ALN cortical node coupled to a thalamic "
    "node. The fitting procedure searches over eight parameters: mue, mui, b, "
    "tauA, g_LK, g_h, c_th2ctx, and c_ctx2th. These include cortical background "
    "inputs and adaptation parameters, thalamic TCR conductances, and directed "
    "thalamocortical / corticothalamic coupling strengths. Each candidate is "
    "simulated for 60 s at 1 kHz; the first 5 s are discarded as burn-in before "
    "any metric is computed. The search uses "
    "scipy.optimize.differential_evolution (strategy='best1bin', population size "
    "20, 30 generations), and all stochastic components (Ornstein–Uhlenbeck "
    "noise and node initialization) are seeded deterministically (seed 42) so "
    "metrics are reproducible across calls with identical parameters. Each "
    "reported version corresponds to a single DE run under one seed; we therefore "
    "treat per-version feasible counts as point estimates without run-to-run "
    "variance and caution against over-reading small differences (Sec. V-B)."
)
h2("C.  V7 fitness and feasibility")
para("The V7 feasible-solution reward is:", space_after=3)
rich([("fitness = 0.50 * shape_r + 0.25 * so_power + 0.25 * spindle_power",
       {"mono": True})], size=9, space_after=6)
para(
    "where shape_r is the FOOOF residual spectral correlation above. Unlike "
    "earlier versions, V7 computes so_power and spindle_power from event-level "
    "quantities:", space_after=3)
rich([("so_power      = clip((T4_q - 1) / 4, 0, 1)", {"mono": True})],
     size=9, space_after=0)
rich([("spindle_power = clip(T12_n_verified / 15, 0, 1)", {"mono": True})],
     size=9, space_after=6)
para(
    "The denominators (4 and 15) are empirical reward-saturation scales, not "
    "normative targets. In particular, 15 is the spindle-event count at which "
    "spindle_power saturates to 1 over the 60 s window; it is set above "
    "realistic event counts to keep the reward discriminating among candidates, "
    "and should be read purely as an internal normalization factor—not a "
    "claim that physiological N3 contains 15 spindles per window (the held-out "
    "density measured here is ≈ 2.4–3.0 events/min, far lower; "
    "Sec. III-E)."
)
para(
    "This creates partial overlap between the reward and constraints, so we do "
    "not treat the V7 fitness score as held-out validation. Instead the paper "
    "uses the version history and feasibility audit as diagnostic evidence and "
    "separately reports held-out checks. Infeasible points are ranked by "
    "continuous constraint satisfaction rather than a single cliff penalty, "
    "preserving a search gradient: candidates that nearly satisfy the audit are "
    "distinguished from those that fail all checks, while all feasible solutions "
    "remain ranked above infeasible ones."
)
h2("D.  Event-level constraints")
para("V7 uses 12 binary constraints in four physiological categories:",
     space_after=3)
for txt in [
    "Cortical bistability: DOWN state, UP state, and sustained UP duration "
    "(T1–T3).",
    "Slow-oscillation structure: SO peak sharpness/frequency and inter-burst-"
    "interval regularity (T4, T6).",
    "Spindle reality: spindle spectral width, envelope burstiness, event "
    "count/duration, and event-internal sigma verification (T5, T7, T8, T12).",
    "SO–spindle coordination: PAC strength, phase preference, and "
    "up/down-ratio / lag-like coupling metric (T9–T11).",
]:
    doc.add_paragraph(txt, style="List Paragraph")
para(
    "The metric forms are physiologically motivated, but the thresholds are "
    "empirical and subject-specific; the paper does not claim every threshold is "
    "literature-derived. Two points guard against the reading that constraints "
    "were simply tuned until V7 passed. First, each constraint is motivated by a "
    "failure observed in an earlier version, independently of whether the final "
    "V7 point satisfies it: the UP-state tests (T1–T3) respond to "
    "persistent-DOWN solutions in V2–V3; the PAC three-pack (T9–T11) "
    "responds to SO–spindle decoupling in V5; and event-internal sigma "
    "verification (T12) responds to fake spindle transients in V6. Each "
    "motivating diagnostic precedes the V7 reward redesign and does not depend on "
    "its outcome. Second, the PAC metrics (T9–T11) use a convention-free, "
    "cycle-by-cycle estimator rather than the bandpass-plus-Hilbert phase used in "
    "early prototypes, which systematically biased the preferred SO phase for "
    "spike-like firing-rate signals; we therefore report PAC strength as a "
    "constraint but avoid preferred-phase point claims."
)
h2("E.  Held-out checks")
para(
    "To test whether event-level feasibility also improved independent "
    "statistics, we implemented two held-out checks, neither of which calls the "
    "V7 constraint function or reuses T1–T12, shape_r, so_power, or "
    "spindle_power."
)
rich([("SO average-cycle morphology. ", {"bold": True}),
      ("Real SC4001 N3 EEG and simulated cortical activity are bandpassed to "
       "the SO range, aligned to SO cycle anchors, normalized, averaged into "
       "templates, and compared by RMSE and correlation.", {})])
rich([("Cortical spindle density. ", {"bold": True}),
      ("Because the SO template primarily probes low-frequency cycle shape and "
       "can look unremarkable even for degenerate event dynamics, we add an "
       "event-level statistic disjoint from the audit: empirical spindle density "
       "(events/min), measured with a standard sigma-band RMS detector "
       "(band-pass to sigma, moving-RMS envelope, threshold = mean + 1.5·SD, "
       "accept events of 0.5–3.0 s). This detector is deliberately disjoint "
       "from the fitness machinery: it operates on the cortical (EEG-analog) "
       "signal, whereas T8/T12 operate on the thalamic signal; it uses a "
       "moving-RMS envelope and mean+SD threshold, whereas T8 uses a "
       "75th-percentile Gaussian-smoothed Hilbert envelope and T12 adds Welch "
       "peak-inside verification; and it uses 0.5–3.0 s duration bounds "
       "rather than T8’s 0.3–2.0 s. We apply it identically to real EEG "
       "and to the simulated cortical firing rate of V1 and V7, over 300 s "
       "simulations and across two sigma bands (10–14 Hz, the model’s "
       "T-current resonance, and 11–15 Hz, the EEG sigma band), so the "
       "result is robust to duration and band choice. Using a held-out statistic "
       "able to favor V7—not only one favoring V1—lets the event-level "
       "result be read in either direction.", {})])

# ============================================================================
# IV. EXPERIMENTS
# ============================================================================
h1("IV.  Experiments")
h2("A.  V1–V7 as an experimental design")
para(
    "The evidence lies in the V1–V7 sequence, where each version exposed a "
    "new objective loophole (Table I).")
for lead, body in [
    ("V1: spectral-only fitting. ",
     "The initial Fig. 7-style baseline used FOOOF residual spectral similarity "
     "and SO/spindle spectral rewards. It reached high spectral similarity "
     "(shape_r = 0.8586) but imposed no event-level guarantees—the cleanest "
     "exhibit for the claim that spectral fit can look good without proving "
     "mechanistic validity."),
    ("V2–V3: persistent-DOWN loophole. ",
     "Weak dynamics checks could be satisfied by solutions with inadequate "
     "UP-state behavior. V3 added explicit UP-state existence and sustained-UP "
     "tests, closing the persistent-DOWN loophole."),
    ("V4: constraints separated from rewards. ",
     "V4 separated feasibility constraints from reward ranking and introduced SO "
     "sharpness, SO regularity, and spindle envelope burstiness checks. It "
     "produced feasible solutions, but only rarely (22/4703)."),
    ("V5: PAC constraints. ",
     "V5 added a three-part PAC audit to prevent SO–spindle decoupling. The "
     "best solution still had spindle_power = 0, showing PAC constraints alone "
     "did not guarantee robust spindle rewards."),
    ("V6: fake-spindle exposure. ",
     "V6 fixed T8 event detection and added T12 to verify that detected events "
     "contain sigma-dominant oscillations. This made the search stricter but "
     "exposed another failure: only 1/4960 evaluations were feasible."),
    ("V7: reward and bound redesign. ",
     "Diagnostics showed FOOOF spindle rewards were effectively unreachable in "
     "SO-stable regimes, 30 s simulations made T6 noisy, T5’s original FWHM "
     "threshold rejected mechanistically valid narrow spindles, and large "
     "c_th2ctx destabilized cortical SO. V7 doubled simulation duration to 60 s, "
     "replaced FOOOF SO/spindle rewards with T4/T12-derived event rewards, "
     "tightened c_th2ctx, and relaxed T5 FWHM to admit narrow T-current "
     "resonance peaks."),
]:
    rich([(lead, {"bold": True}), (body, {})])
h2("B.  Evaluation measures")
para(
    "We report: (1) best score and shape_r for each version; (2) feasible rate "
    "for versions with explicit feasibility rules; (3) V7 final T1–T12 "
    "metrics; and (4) held-out SO waveform morphology distance and cortical "
    "spindle density. The held-out results are not used to select V7; they test "
    "and constrain the interpretation of V7’s event-level improvements."
)

# ============================================================================
# V. RESULTS
# ============================================================================
h1("V.  Results")
h2("A.  Spectral fit alone can look successful")
para(
    "The V1 spectral-only baseline achieved a high best shape_r = 0.8586, "
    "confirming the optimizer can find a solution whose FOOOF residual spectrum "
    "resembles the target EEG. V1 did not enforce UP/DOWN alternation, spindle "
    "event structure, event-internal sigma verification, or SO–spindle "
    "coupling."
)
para(
    "To convert this from “unchecked” to “demonstrably "
    "invalid,” we evaluated the canonical V1 selected parameter vector (the "
    "shape_r = 0.8586 point) through the same event-level audit later formalized "
    "as V7’s T1–T12, using identical 60 s simulation, 5 s burn-in, and "
    "seed (42). The V1 point passes only 5/12 constraints. It retains a DOWN "
    "state (T1) and a momentary UP excursion (T2, max cortical rate 38.6 Hz), but "
    "fails: T3—longest sustained UP state 86 ms, below the 100 ms threshold, "
    "so UP/DOWN alternation is spiky rather than bistable; T4—SO spectral "
    "peak Q = 1.68 (< 2.0), no sharp slow-oscillation peak; T6—inter-burst-"
    "interval CV = 1.03 (≫ 0.4), an irregular slow rhythm; T7—spindle-"
    "envelope CV = 0.52 (< 0.7), not waxing–waning; and T9–T11—"
    "modulation index exactly 0, a complete absence of SO–spindle coupling. "
    "The corresponding cortical time series is shown in Fig. 3. This is the "
    "paper’s existence object: a single solution spectrally similar to the "
    "target (shape_r = 0.8586) yet absent two core N3 phenomena—sustained UP "
    "states and SO–spindle coupling. Re-scoring V1 under V7’s audit adds "
    "no degrees of freedom (the harness reproduces V7’s stored 12/12 "
    "exactly, confirming fidelity). Thus V1 is not merely unvalidated—under "
    "explicit audit it is invalid on named criteria—while remaining "
    "spectrally convincing. For reference, the V3 dynamics point also passes 5/12 "
    "(failing T4, T6, T7, T9–T11, T12), so the spectral-versus-event gap is "
    "not unique to V1."
)
h2("B.  Feasibility changed non-monotonically across versions")
add_table(
    ["Version", "Best score", "Best shape_r", "Feasible rate", "Interpretation"],
    [
        ["V1", "0.8586", "0.8586", "N/A",
         "Spectral objective can look good without event guarantees"],
        ["V3", "0.7722", "0.8334", "N/A",
         "UP-state tests close persistent-DOWN loophole"],
        ["V4", "0.4481", "0.5449", "22/4703",
         "Event feasibility introduced but rare"],
        ["V5", "0.5296", "0.5867", "5/4960",
         "PAC constraints added; spindle reward still problematic"],
        ["V6", "0.3080", "0.6159", "1/4960", "Fake-spindle loophole exposed"],
        ["V7", "0.7759", "0.6226", "364/4960",
         "Event-level loopholes closed much more often"],
    ],
    "Table II.  Version-level best score, shape_r, and feasible rate.",
)
para(
    "The non-monotonic pattern matters: V5 and V6 did not simply improve over "
    "earlier versions; they became stricter and revealed new failures. V7 "
    "recovered feasibility only after reward targets and bounds were redesigned "
    "according to those diagnostics."
)
para(
    "Two caveats keep this trajectory honest. First, the feasible rate is not a "
    "clean measure of loophole closure, because between V6 and V7 four things "
    "changed at once: simulation duration (30 → 60 s, reducing T6 sampling "
    "noise), the c_th2ctx bound (tightened), the SO/spindle reward definitions "
    "(FOOOF-based → reachable event-based), and the T5 spindle-width "
    "threshold (relaxed from > 2.0 Hz to > 0.2 Hz, mechanically admitting narrow "
    "T-current resonance peaks previously rejected). The last loosens a gate, so "
    "part of the recovery is threshold relaxation rather than better physiology. "
    "Re-scoring the V7 evaluation population under V6’s original T5 threshold "
    "isolates this: of V7’s 364 feasible points, only 156/4960 remain "
    "feasible under V6’s strict FWHM > 2.0 Hz rule, so 208 (57%) of the "
    "recovery is attributable to T5 relaxation alone and 156 to the "
    "duration/reward/bound redesign. The V7 selected solution itself has spindle "
    "FWHM exactly 2.0 Hz: under V6’s strict threshold it would be "
    "infeasible, confirming T5 is load-bearing rather than incidental. Second, "
    "because each version is a single seeded DE run, small absolute counts are "
    "within plausible run-to-run variation; the V5-vs-V6 difference (5 vs 1) "
    "should not be read as meaningful, and we draw conclusions only from the "
    "order-of-magnitude V6 → V7 shift after de-confounding."
)
h2("C.  V7 selected point passed the event-level audit")
para(
    "The selected V7 point passed all 12 constraints (n_passed = 12) with "
    "score = 0.7759, shape_r = 0.6226, T4_q = 4.433, T4_freq = 1.25 Hz, "
    "T6_ibi_cv = 0.197, T8_n_sp_events = 25, and T12_n_verified = 20."
)
para(
    "This requires careful interpretation. V7’s shape_r is lower than "
    "V1’s spectral-only score, so V7 is not a better spectral fit. The "
    "comparison is also not on identical ground: V1 searched broad bidirectional "
    "coupling bounds, whereas V7 tightened c_th2ctx for SO stability, so part of "
    "the shape_r gap reflects a smaller, more physiologically constrained search "
    "space rather than a pure spectral/event trade-off. With that caveat, "
    "V7’s advantage is event-level feasibility: the final point satisfies "
    "the explicit audit for cortical bistability, SO regularity, spindle event "
    "structure, and SO–spindle coordination."
)
h2("D.  Held-out checks give mixed evidence")
para(
    "We evaluated two held-out statistics, both on the canonical V1 point "
    "(shape_r = 0.8586)."
)
rich([("(1) SO average-cycle waveform morphology. ", {"bold": True}),
      ("Using 142 clean SC4001 N3 epochs from EEG Fpz-Cz, the real template "
       "comprised 3881 SO snippets. Distances to the real template are given in "
       "Table III. On this metric the canonical V1 point is the worst of the "
       "three, and V7 is closer to real EEG than V1 (RMSE 0.39 vs 0.57).", {})])
add_table(
    ["Version", "RMSE", "Correlation"],
    [
        ["V1 canonical (shape_r = 0.8586)", "0.5685", "0.8390"],
        ["V3 dynamics", "0.3424", "0.9414"],
        ["V7 event-constrained", "0.3909", "0.9234"],
    ],
    "Table III.  Held-out SO average-cycle waveform distances to the real SC4001 "
    "template (142 clean N3 epochs; 3881 SO snippets).",
)
rich([("(2) Cortical spindle density. ", {"bold": True}),
      ("Real SC4001 N3 EEG has a cortical spindle density of 2.41/min "
       "(10–14 Hz) to 2.96/min (11–15 Hz); the canonical V1 point "
       "reaches 1.22–1.42/min—about half the real value—while V7 "
       "produces zero cortical spindles in both bands over a 300 s simulation. On "
       "this metric V1 is closer to real EEG than V7.", {})])
para(
    "The two checks disagree: V7 is closer on SO morphology, V1 is closer on "
    "spindle density. Neither is decisive. The SO average-cycle metric is "
    "narrow—it normalizes amplitude and probes only low-frequency cycle "
    "shape—so V7’s edge there does not establish global fidelity; and "
    "V7’s complete absence of cortical spindles is a real deficiency the "
    "audit did not prevent (Sec. VI-F). We therefore retain only the "
    "audited-invalidity claim (Sec. V-A) and disclaim any global ranking. This is "
    "why the result settles on “invalid” (under explicit audit) rather "
    "than “wrong”: the audit demonstrates that the spectrally "
    "convincing V1 point violates named event-level criteria, but the held-out "
    "evidence does not support the stronger claim that either model is the more "
    "faithful one. V1 is not thereby validated—it passes only 5/12 audit "
    "checks and under-produces cortical spindles by about half—and V7, "
    "though it passes all 12, fails the cortical spindle check. Neither solution "
    "is correct; the spectral objective and the event audit each select a "
    "differently-deficient model."
)

# ============================================================================
# VI. DISCUSSION
# ============================================================================
h1("VI.  Discussion and Limitations")
for head, body in [
    ("A.  Main interpretation",
     "The central result is not that V7 is globally superior to V1. It is that "
     "V1 demonstrates the danger of spectral fitting: a high residual spectral "
     "score can coexist with missing event-level guarantees. The V1–V7 "
     "sequence then shows how the optimizer exploited each objective or "
     "constraint loophole until a more explicit event-level audit was "
     "introduced. Under the adversarial-optimizer framing, the post-hoc nature "
     "of several constraints is not something to hide—it is evidence: each "
     "constraint was introduced because an apparently reasonable previous "
     "objective admitted an unintended solution. The result is a diagnostic "
     "methodology for discovering objective failures in personalized neural mass "
     "fitting."),
    ("B.  Single-subject scope",
     "This is a single-subject case study. It demonstrates that the failure mode "
     "exists for SC4001 N3, not that the same thresholds or final parameters "
     "generalize across subjects. A population-level analysis would require "
     "repeating the procedure across subjects and assessing whether the same "
     "loopholes and constraint responses recur."),
    ("C.  Empirical thresholds",
     "The constraints combine physiologically interpretable metrics with "
     "empirical thresholds. UP/DOWN states, spindle event duration, spindle "
     "envelope burstiness, and PAC are grounded in sleep physiology, but the "
     "exact values used here are subject- and model-specific. The manuscript "
     "should therefore be read as a diagnostic study of objective failure, not "
     "as a universal definition of N3 physiology."),
    ("D.  Partial circularity",
     "V7’s feasible-solution fitness includes a T4-derived SO reward and a "
     "T12-derived spindle reward, creating overlap between fitness and "
     "constraints. We therefore do not use the V7 score itself as held-out "
     "validation. Instead we report the overlap explicitly and include two "
     "independent held-out checks (Sec. V-D), which give mixed evidence and do "
     "not let us rank V1 against V7."),
    ("E.  Mixed held-out results",
     "The two held-out checks narrow the conclusion, and they disagree. On SO "
     "average-cycle morphology, V7 is closer to real EEG than canonical V1 (RMSE "
     "0.39 vs 0.57); on cortical spindle density, V1 is closer because V7 "
     "produces no cortical spindles at all. Together they prevent both the claim "
     "that V7 is simply “more EEG-like” and the opposite claim that V1 "
     "is more EEG-like—the held-out evidence does not rank the two. A model "
     "can pass the event audit while still being deficient on an independent "
     "statistic (here, cortical spindle expression), and a spectrally selected "
     "model can be closer on one waveform statistic yet invalid under the audit. "
     "Neither model is correct: V1 fails 5/12 audit checks and under-produces "
     "cortical spindles, and V7, though it passes all 12, produces zero cortical "
     "spindles. The contribution is diagnostic, not a ranking."),
]:
    h2(head)
    para(body)
h2("F.  Residual loophole: thalamic spindles without cortical expression")
para(
    "The spindle-density check exposes a residual loophole inside the "
    "“closed” V7 audit, and it is the sharpest illustration of this "
    "paper’s thesis. V7 passes T8 (25 detected spindle events) and T12 (20 "
    "sigma-verified events)—but those constraints are evaluated on the "
    "thalamic firing rate. Measured on the cortical signal, the EEG-analog "
    "quantity, V7 yields exactly zero events over 300 s in both the 10–14 Hz "
    "and 11–15 Hz bands. The mechanism is direct: V7’s selected "
    "c_th2ctx = 0.0127 is roughly eight times smaller than V1’s 0.099, "
    "because V7 tightened thalamocortical coupling toward zero for SO stability. "
    "At that coupling, thalamic spindles barely reach cortex, so the events "
    "satisfying the audit would not appear in EEG."
)
para(
    "This is specification gaming one level deeper than the failures V7 was "
    "designed to close. The V7 audit verified spindle reality (an event must "
    "contain a genuine sigma oscillation, T12) but not spindle observability in "
    "the signal the EEG actually records. The optimizer, rewarded for thalamic "
    "spindle events, found a regime that produces them on the unobserved node "
    "while the SO-stability bound suppressed their cortical expression. The fix "
    "is straightforward to state—verify spindles on the cortical signal, or "
    "constrain c_th2ctx from below—but the loophole was invisible until an "
    "independent, signal-matched held-out check revealed it. It also reinforces "
    "that V1 is no remedy: V1’s larger coupling yields some cortical "
    "spindles (1.2–1.4/min) but still only about half the real density "
    "(2.4–3.0/min), and V1 fails 5/12 audit checks. Both models are "
    "deficient in different ways; the audit and the spectrum each hide a "
    "different failure."
)
h2("G.  Stage 2 SBI")
para(
    "Simulation-based inference is reserved for a future journal extension "
    "focused on posterior identifiability. It addresses a different "
    "question—which parameters are identifiable under selected summary "
    "statistics—and is not used here to prove V7 superiority. This paper "
    "remains focused on Stage 1 objective diagnosis and event-level loophole "
    "closure."
)

# ============================================================================
# VII. CONCLUSION
# ============================================================================
h1("VII.  Conclusion")
para(
    "This single-subject diagnostic study shows that spectral objectives can "
    "select thalamocortical model parameters that are shape-alike yet leave the "
    "intended dynamics unverified—and that, when those dynamics are "
    "explicitly audited, the spectrally convincing V1 solution is invalid on "
    "named event-level criteria (it passes only 5 of 12 checks, with no "
    "sustained UP states and zero SO–spindle coupling). Treating the "
    "optimizer as a loophole-searcher revealed a sequence of failure modes "
    "across V1–V6, including persistent-DOWN behavior, SO–spindle "
    "decoupling, fake spindle events, and unreachable spectral spindle rewards. "
    "V7 closed these event-level loopholes and increased feasible solutions from "
    "1/4960 in V6 to 364/4960 (in part through threshold relaxation we quantify "
    "separately), yielding a final point that passed all 12 event-level "
    "constraints. At the same time, two independent held-out checks gave mixed "
    "evidence—V7 closer on SO waveform morphology, V1 closer on cortical "
    "spindle density—and a residual loophole emerged: V7 passes its thalamic "
    "spindle audit yet expresses zero spindles in the EEG-analog cortical signal "
    "at its small thalamocortical coupling. We therefore make no claim that V7 is "
    "the better model; V1 is itself invalid under the audit (5/12) and "
    "under-produces cortical spindles by roughly half. Both solutions are "
    "deficient in different ways. The contribution is not a universal fitting "
    "method, nor a ranking of V1 against V7, but a transferable diagnostic "
    "workflow: anchor a spectral objective, treat the optimizer as an adversary, "
    "and use each feasibility collapse—and each independent held-out "
    "check—to localize the next loophole, including loopholes that survive "
    "inside an apparently closed audit. We report this for one subject; whether "
    "the same loopholes recur across subjects is the natural next study."
)

# ============================================================================
# TABLE I (loophole map) — placed near end like the source figure plan
# ============================================================================
add_table(
    ["Ver.", "Loophole exposed", "Constraint / change introduced"],
    [
        ["V1", "Spectral fit without event guarantees",
         "Baseline (FOOOF residual + spectral SO/spindle rewards)"],
        ["V2–V3", "Persistent-DOWN solutions",
         "UP-state existence + sustained-UP tests (T1–T3)"],
        ["V4", "Rewards entangled with feasibility",
         "Constraints separated; SO sharpness/regularity, spindle burstiness"],
        ["V5", "SO–spindle decoupling", "PAC three-pack (T9–T11)"],
        ["V6", "Fake spindle transients",
         "Event-internal sigma verification (T12)"],
        ["V7",
         "Unreachable FOOOF rewards; noisy T6; over-strict T5; large c_th2ctx",
         "60 s sims; T4/T12 event rewards; tightened c_th2ctx; relaxed T5"],
    ],
    "Table I.  Loophole-to-constraint map across V1–V7.",
)

# ---- Figures ---------------------------------------------------------------
h2("Figures")

# Fig. 1: spectral-fit loophole schematic (generated PNG).
add_figure(FIG1, width_in=5.5)
fig_caption(
    "Fig. 1.  Schematic of the spectral-fit loophole. Two models can share a "
    "nearly identical power spectrum (left; high shape_r) while differing sharply "
    "in time-domain dynamics (right): Model A shows sustained UP/DOWN bistability "
    "with coupled spindles, whereas Model B produces spiky transients with no "
    "sustained UP state and no SO–spindle coupling. The spectral objective cannot "
    "separate them (traces are illustrative).")

# Fig. 2: feasible-rate recovery (real PNG).
add_figure(FIG2, width_in=5.0)
fig_caption(
    "Fig. 2.  Feasible-solution recovery (V4–V7), log scale. Feasible counts "
    "fall from 22/4703 (V4) to 1/4960 (V6) as constraints tighten, then jump to "
    "364/4960 (V7) after the reward/bound redesign. The hatched overlay marks the "
    "156/4960 V7 points that remain feasible under V6’s strict FWHM > 2.0 Hz "
    "rule—i.e. about half the recovery is T5 relaxation rather than loophole "
    "closure (Sec. V-B).")

# Fig. 3: existence object (real PNG).
add_figure(FIG3, width_in=5.5)
fig_caption(
    "Fig. 3.  Existence object—cortical excitatory firing rate of the V1, "
    "V3, and V7 selected points under the same T1–T12 audit (60 s sim, 5 s "
    "burn-in, seed 42). The canonical V1 point (shape_r = 0.8586) passes only "
    "5/12 and shows spiky UP excursions with no SO–spindle coupling, whereas "
    "V7 passes 12/12 with sustained UP/DOWN alternation; each panel annotates the "
    "failed checks (Sec. V-A).")

# Fig. 4: held-out mixed evidence — two real PNG panels.
add_figure(FIG4A, width_in=5.0)
para("(a)", align=WD_ALIGN_PARAGRAPH.CENTER, size=9, space_after=4)
add_figure(FIG4B, width_in=5.0)
para("(b)", align=WD_ALIGN_PARAGRAPH.CENTER, size=9, space_after=3)
fig_caption(
    "Fig. 4.  Held-out checks giving mixed evidence (canonical V1 throughout). "
    "(a) SO average-cycle waveform distance to the real EEG template (RMSE and "
    "1 − corr): V1 is the worst (RMSE 0.57), V7 (0.39) beats V1. "
    "(b) Cortical spindle density over 300 s in two sigma bands: real EEG "
    "≈ 2.4–3.0/min, V1 ≈ 1.2–1.4/min, and V7 = 0/min in both "
    "10–14 and 11–15 Hz bands. The two checks disagree, so neither model is "
    "ranked best overall.")

# Fig. 5: residual-loophole schematic (generated PNG).
add_figure(FIG5, width_in=5.5)
fig_caption(
    "Fig. 5.  Residual-loophole schematic. V7’s audit (T8/T12) verifies spindle "
    "events on the thalamic node, but the EEG-analog signal is the cortical node. "
    "At V7’s small coupling (c_th2ctx ≈ 0.0127) thalamic spindles do not "
    "propagate, so cortical spindle density is ≈ 0 even though the audit "
    "passes—the reality check and the held-out measurement inspect different "
    "nodes.")

# ============================================================================
# REFERENCES
# ============================================================================
h1("References")
refs = [
    "C. Cakan, N. Jajcay, and K. Obermayer, “neurolib: A Simulation "
    "Framework for Whole-Brain Neural Mass Modeling,” Cognitive Computation, "
    "vol. 15, pp. 1132–1152, 2023, doi:10.1007/s12559-021-09931-9.",
    "N. Jajcay, C. Cakan, and K. Obermayer, “Cross-Frequency Slow "
    "Oscillation–Spindle Coupling in a Biophysically Realistic "
    "Thalamocortical Neural Mass Model,” Front. Comput. Neurosci., vol. 16, "
    "art. 769860, 2022, doi:10.3389/fncom.2022.769860.",
    "B. Kemp, A. H. Zwinderman, B. Tuk, H. A. C. Kamphuisen, and J. J. L. "
    "Oberye, “Analysis of a sleep-dependent neuronal feedback loop: the "
    "slow-wave microcontinuity of the EEG,” IEEE Trans. Biomed. Eng., "
    "vol. 47, no. 9, pp. 1185–1194, 2000, doi:10.1109/10.867928. Database: "
    "A. L. Goldberger et al., “PhysioBank, PhysioToolkit, and PhysioNet,"
    "” Circulation, vol. 101, no. 23, pp. e215–e220, 2000, "
    "doi:10.1161/01.cir.101.23.e215.",
    "P. A. Robinson, C. J. Rennie, and D. L. Rowe, “Dynamics of large-scale "
    "brain activity in normal arousal states and epileptic seizures,” Phys. "
    "Rev. E, vol. 65, no. 4, art. 041924, 2002, doi:10.1103/PhysRevE.65.041924.",
    "M. Steriade, A. Nuñez, and F. Amzica, “A novel slow (<1 Hz) "
    "oscillation of neocortical neurons in vivo: depolarizing and hyperpolarizing "
    "components,” J. Neurosci., vol. 13, no. 8, pp. 3252–3265, 1993, "
    "doi:10.1523/JNEUROSCI.13-08-03252.1993.",
    "M. Massimini, R. Huber, F. Ferrarelli, S. Hill, and G. Tononi, “The "
    "sleep slow oscillation as a traveling wave,” J. Neurosci., vol. 24, "
    "no. 31, pp. 6862–6870, 2004, doi:10.1523/JNEUROSCI.1318-04.2004.",
    "M. Mölle, T. O. Bergmann, L. Marshall, and J. Born, “Fast and slow "
    "spindles during the sleep slow oscillation: disparate coalescence and "
    "engagement in memory processing,” Sleep, vol. 34, no. 10, "
    "pp. 1411–1421, 2011, doi:10.5665/SLEEP.1290.",
    "R. F. Helfrich, B. A. Mander, W. J. Jagust, R. T. Knight, and M. P. Walker, "
    "“Old brains come uncoupled in sleep: slow wave–spindle synchrony, "
    "brain atrophy, and forgetting,” Neuron, vol. 97, no. 1, "
    "pp. 221–230.e4, 2018, doi:10.1016/j.neuron.2017.11.020.",
    "A. B. L. Tort, R. Komorowski, H. Eichenbaum, and N. Kopell, “Measuring "
    "phase-amplitude coupling between neuronal oscillations of different "
    "frequencies,” J. Neurophysiol., vol. 104, no. 2, pp. 1195–1210, "
    "2010, doi:10.1152/jn.00106.2010.",
    "R. T. Canolty et al., “High gamma power is phase-locked to theta "
    "oscillations in human neocortex,” Science, vol. 313, no. 5793, "
    "pp. 1626–1628, 2006, doi:10.1126/science.1128115.",
    "B. A. Mander et al., “β-amyloid disrupts human NREM slow waves and "
    "related hippocampus-dependent memory consolidation,” Nat. Neurosci., "
    "vol. 18, no. 7, pp. 1051–1057, 2015, doi:10.1038/nn.4035.",
    "F. Ferrarelli, R. Huber, M. J. Peterson, M. Massimini, M. Murphy, B. A. "
    "Riedner, A. Watson, P. Bria, and G. Tononi, “Reduced sleep spindle "
    "activity in schizophrenia patients,” Am. J. Psychiatry, vol. 164, "
    "no. 3, pp. 483–492, 2007, doi:10.1176/ajp.2007.164.3.483.",
    "R. Vallat and M. P. Walker, “An open-source, high-performance tool for "
    "automated sleep staging,” eLife, vol. 10, art. e70092, 2021, "
    "doi:10.7554/eLife.70092.",
    "V. Krakovna, J. Uesato, V. Mikulik, M. Rahtz, T. Everitt, R. Kumar, "
    "Z. Kenton, J. Leike, and S. Legg, “Specification gaming: the flip side "
    "of AI ingenuity,” DeepMind Safety Research, 2020.",
    "D. Amodei, C. Olah, J. Steinhardt, P. Christiano, J. Schulman, and "
    "D. Mané, “Concrete Problems in AI Safety,” arXiv:1606.06565, "
    "2016; J. Skalse, N. H. R. Howe, D. Krasheninnikov, and D. Krueger, "
    "“Defining and Characterizing Reward Hacking,” in Adv. Neural Inf. "
    "Process. Syst. (NeurIPS), 2022, arXiv:2209.13085.",
    "T. Donoghue, M. Haller, E. J. Peterson, P. Varma, P. Sebastian, R. Gao, "
    "T. Noto, A. H. Lara, J. D. Wallis, R. T. Knight, A. Shestyuk, and "
    "B. Voytek, “Parameterizing neural power spectra into periodic and "
    "aperiodic components,” Nat. Neurosci., vol. 23, pp. 1655–1665, "
    "2020, doi:10.1038/s41593-020-00744-x.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}]  {r}")
    run.font.name = FONT
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(3)

doc.save(str(OUT))
print(f"Saved: {OUT}")
print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
